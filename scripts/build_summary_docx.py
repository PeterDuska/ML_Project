from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_paragraphs(doc: Document, text: str) -> None:
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        doc.add_paragraph(line)


def main() -> None:
    src = Path("summary.md")
    if not src.is_file():
        raise SystemExit("summary.md not found")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    current_para: list[str] = []
    for raw in src.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            if current_para:
                add_paragraphs(doc, "\n".join(current_para))
                current_para = []
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            if current_para:
                add_paragraphs(doc, "\n".join(current_para))
                current_para = []
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if not line:
            if current_para:
                add_paragraphs(doc, "\n".join(current_para))
                current_para = []
            continue
        current_para.append(line)

    if current_para:
        add_paragraphs(doc, "\n".join(current_para))

    doc.save("summary.docx")


if __name__ == "__main__":
    main()
